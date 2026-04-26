import os
from docx import Document

# 你的目标文件夹（直接用你给的路径，不用改！）
folder_path = r"C:\Users\未来可期\Desktop\文档\从官方文档 - 副本"

# 遍历文件夹里所有文件
for filename in os.listdir(folder_path):
    # 只处理 .docx 文件
    if filename.endswith(".docx"):
        # 拼接完整文件路径
        docx_path = os.path.join(folder_path, filename)
        
        try:
            # 打开docx
            doc = Document(docx_path)
            full_text = []

            # 提取正文
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text.append(txt)

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))

            # 生成 txt 文件名（原名字 + .txt）
            txt_filename = os.path.splitext(filename)[0] + ".txt"
            txt_path = os.path.join(folder_path, txt_filename)

            # 写入文件
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write("\n".join(full_text))

            print(f"✅ 已转换：{filename} → {txt_filename}")
        
        except Exception as e:
            print(f"❌ 处理失败 {filename}：{str(e)}")

print("\n🎉 全部处理完成！")